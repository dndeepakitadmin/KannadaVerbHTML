import streamlit as st
from io import BytesIO
import zipfile

# ✅ Optional DOCX library (auto handles missing one)
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ✅ List of 50 Kannada verbs
verbs = [
    "ತಿನ್ನು", "ಕುಡಿ", "ಓದು", "ಬರೆಯು", "ಹೋಗು", "ಬಾ", "ನೋಡು", "ಕೇಳು", "ಮಲಗು", "ನಿಲ್ಲು",
    "ಕುಳಿತು", "ಕೂರು", "ನಗು", "ಅಳು", "ನಡೆಯು", "ಓಡು", "ಎದ್ದು", "ತಗೆದುಕೋ", "ಕೊಡು", "ಇಡು",
    "ತೆರೆ", "ಮುಚ್ಚು", "ತೋರಿಸು", "ಹೇಳು", "ಕೇಳು", "ಮಾಡು", "ಇರು", "ಬಸು", "ತಗೆ", "ಬೀಳು",
    "ನಿಲ್ಲಿಸು", "ಹರಿ", "ಆಡು", "ಕಲಿತುಕೋ", "ನೆನಪುಮಾಡು", "ಮರೆತುಬಿಡು", "ನೋಡಿಕೊಳ್ಳು",
    "ಹೋಗಿಬಿಡು", "ತರಲು", "ಕತ್ತರಿಸು", "ತೊಳೆಯು", "ತಿನ್ನಿಸು", "ಹೊಡೆ", "ಹಿಡಿ",
    "ಓದಿ", "ಹಾಡು", "ನೃತ್ಯಮಾಡು", "ನಿಲ್ಲಿಸು", "ಕಳೆ", "ಗೆಲ್ಲು"
]

# ✅ Function to generate Kannada tense sentences (placeholder structure)
def generate_html(verb):
    html_content = f"""
    <html>
    <head>
      <meta charset="UTF-8">
      <style>
        body {{
          font-family: 'Noto Sans Kannada', sans-serif;
          background-color: #f9f9f9;
          padding: 20px;
          line-height: 1.8;
          font-size: 20px;
        }}
        h2 {{
          color: #2b4b7c;
        }}
        .person {{
          margin-top: 25px;
          background: #fff;
          border-radius: 10px;
          padding: 15px;
          box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
      </style>
    </head>
    <body>
      <h2>ಕ್ರಿಯಾಪದ: {verb}</h2>
      <div class="person">
        🧍 ನಾನು <br>
        ಹಿಂದಿನ ಕಾಲ: ನಾನು ನಿನ್ನೆ {verb}ದೆಯೆ. <br>
        → ಮೈ ಕಲ್ {verb} ಕಿಯಾ. <br><br>
        ವರ್ತಮಾನ ಕಾಲ: ನಾನು ಈಗ {verb}ತ್ತಿದ್ದೇನೆ. <br>
        → ಮೈ ಅಭಿ {verb} ರಹಾ ಹೂಂ. <br><br>
        ಭವಿಷ್ಯತ್ತಿನ ಕಾಲ: ನಾನು ನಾಳೆ {verb}ತ್ತೇನೆ. <br>
        → ಮೈ ಕಲ {verb} ಕರೂಂಗಾ. <br>
      </div>
    </body>
    </html>
    """
    return html_content

# ✅ Streamlit UI
st.title("📘 ಕನ್ನಡ ಕ್ರಿಯಾಪದ ಪಾಠಗಳು - HTML & DOCX ಡೌನ್‌ಲೋಡ್")

st.write("👇 ಕೆಳಗಿನ ಬಟನ್ ಒತ್ತಿ 50 ಕ್ರಿಯಾಪದಗಳ HTML ಫೈಲುಗಳನ್ನು ZIP ಆಗಿ ಡೌನ್‌ಲೋಡ್ ಮಾಡಿರಿ")

# ✅ Generate ZIP of HTMLs in-memory
zip_buffer = BytesIO()
with zipfile.ZipFile(zip_buffer, "w") as zf:
    for verb in verbs:
        html_data = generate_html(verb)
        zf.writestr(f"{verb}.html", html_data)

zip_buffer.seek(0)

# ✅ HTML ZIP download button
st.download_button(
    label="⬇️ Download All 50 HTML Files (ZIP)",
    data=zip_buffer,
    file_name="Kannada_Verbs_HTML.zip",
    mime="application/zip"
)

# ✅ Optional DOCX download (if installed)
if DOCX_AVAILABLE:
    st.markdown("---")
    st.subheader("📄 DOCX Version (Optional)")
    doc = Document()
    doc.add_heading("ಕನ್ನಡ ಕ್ರಿಯಾಪದ ಪಾಠಗಳು", level=1)
    for verb in verbs:
        doc.add_heading(verb, level=2)
        doc.add_paragraph(f"ನಾನು ನಿನ್ನೆ {verb}ದೆಯೆ.")
        doc.add_paragraph(f"ನಾನು ಈಗ {verb}ತ್ತಿದ್ದೇನೆ.")
        doc.add_paragraph(f"ನಾನು ನಾಳೆ {verb}ತ್ತೇನೆ.")
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    st.download_button(
        label="⬇️ Download DOCX Version",
        data=doc_buffer,
        file_name="Kannada_Verbs.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.warning("⚠️ DOCX library not found. Add `python-docx` to requirements.txt to enable DOCX export.")
